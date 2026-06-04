"""
Document ingestion pipeline: load → parse → chunk → embed → index into Qdrant + ES.
"""
import asyncio
import hashlib
import logging
import uuid
from datetime import datetime, timezone
from typing import Optional

from qdrant_client import AsyncQdrantClient
from qdrant_client.models import Distance, PointStruct, VectorParams

from backend.core.config import settings
from backend.core.database import get_db_context
from backend.services.ai_service.llm.model_router import llm

logger = logging.getLogger(__name__)

CHUNK_SIZE = 800       # tokens approx
CHUNK_OVERLAP = 150
COLLECTION_NAME = "knowledge_base"
VECTOR_DIM = 768       # nomic-embed-text-v2


class IngestionPipeline:
    def __init__(self):
        self._qdrant: Optional[AsyncQdrantClient] = None

    @property
    def qdrant(self) -> AsyncQdrantClient:
        if self._qdrant is None:
            self._qdrant = AsyncQdrantClient(
                host=settings.QDRANT_HOST,
                port=settings.QDRANT_PORT,
                api_key=settings.QDRANT_API_KEY or None,
            )
        return self._qdrant

    async def ensure_collection(self) -> None:
        collections = await self.qdrant.get_collections()
        names = [c.name for c in collections.collections]
        if COLLECTION_NAME not in names:
            await self.qdrant.create_collection(
                collection_name=COLLECTION_NAME,
                vectors_config=VectorParams(size=VECTOR_DIM, distance=Distance.COSINE),
            )
            logger.info(f"Created Qdrant collection: {COLLECTION_NAME}")

    async def ingest_file(
        self,
        content: bytes,
        filename: str,
        title: str,
        source_type: str,
        tenant_id: str,
    ) -> str:
        job_id = str(uuid.uuid4())
        asyncio.create_task(
            self._process_document(content, filename, title, source_type, tenant_id, job_id)
        )
        return job_id

    async def _process_document(
        self,
        content: bytes,
        filename: str,
        title: str,
        source_type: str,
        tenant_id: str,
        job_id: str,
    ) -> None:
        doc_id = str(uuid.uuid4())
        try:
            await self.ensure_collection()

            # Parse text from file
            text = self._extract_text(content, filename)
            if not text.strip():
                logger.warning(f"Empty document: {title}")
                return

            file_hash = hashlib.sha256(content).hexdigest()

            # Create DB record
            async with get_db_context() as db:
                from backend.shared.models.knowledge import KnowledgeDocument
                doc = KnowledgeDocument(
                    id=uuid.UUID(doc_id),
                    tenant_id=uuid.UUID(tenant_id),
                    title=title,
                    source_type=source_type,
                    file_hash=file_hash,
                    status="processing",
                    embedding_model=settings.LLM_EMBEDDING_MODEL,
                )
                db.add(doc)
                await db.commit()

            # Chunk the document
            chunks = self._chunk_text(text, CHUNK_SIZE, CHUNK_OVERLAP)
            logger.info(f"Ingesting '{title}': {len(chunks)} chunks")

            # Embed in batches of 20
            points = []
            batch_size = 20
            for i in range(0, len(chunks), batch_size):
                batch = chunks[i:i + batch_size]
                embeddings = await llm.embed(batch)
                for j, (chunk_text, embedding) in enumerate(zip(batch, embeddings)):
                    chunk_idx = i + j
                    points.append(PointStruct(
                        id=str(uuid.uuid4()),
                        vector=embedding,
                        payload={
                            "tenant_id": tenant_id,
                            "doc_id": doc_id,
                            "title": title,
                            "source_type": source_type,
                            "text": chunk_text,
                            "chunk_index": chunk_idx,
                            "total_chunks": len(chunks),
                            "created_at": datetime.now(timezone.utc).isoformat(),
                        },
                    ))

            # Upsert to Qdrant
            await self.qdrant.upsert(collection_name=COLLECTION_NAME, points=points)

            # Update DB record
            async with get_db_context() as db:
                from backend.shared.models.knowledge import KnowledgeDocument
                from sqlalchemy import update
                stmt = (
                    update(KnowledgeDocument)
                    .where(KnowledgeDocument.id == uuid.UUID(doc_id))
                    .values(
                        status="indexed",
                        chunk_count=len(chunks),
                        indexed_at=datetime.now(timezone.utc),
                    )
                )
                await db.execute(stmt)
                await db.commit()

            logger.info(f"Ingested '{title}': {len(points)} vectors indexed")

        except Exception as exc:
            logger.error(f"Ingestion failed for '{title}': {exc}")
            async with get_db_context() as db:
                from backend.shared.models.knowledge import KnowledgeDocument
                from sqlalchemy import update
                await db.execute(
                    update(KnowledgeDocument)
                    .where(KnowledgeDocument.id == uuid.UUID(doc_id))
                    .values(status="failed")
                )
                await db.commit()

    def _extract_text(self, content: bytes, filename: str) -> str:
        ext = filename.lower().split(".")[-1] if "." in filename else "txt"

        if ext == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(__import__("io").BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        elif ext in ("docx", "doc"):
            from docx import Document
            doc = Document(__import__("io").BytesIO(content))
            return "\n".join(para.text for para in doc.paragraphs)

        elif ext == "md":
            import markdown
            import re
            html = markdown.markdown(content.decode("utf-8", errors="ignore"))
            return re.sub("<[^>]+>", " ", html)

        else:
            return content.decode("utf-8", errors="ignore")

    def _chunk_text(self, text: str, chunk_size: int, overlap: int) -> list[str]:
        """Sentence-aware text chunking."""
        import re
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current = []
        current_len = 0

        for sentence in sentences:
            sentence_len = len(sentence.split())
            if current_len + sentence_len > chunk_size and current:
                chunks.append(" ".join(current))
                # Keep overlap: take last N words from current
                overlap_words = " ".join(current).split()[-overlap:]
                current = overlap_words + sentence.split()
                current_len = len(current)
            else:
                current.extend(sentence.split())
                current_len += sentence_len

        if current:
            chunks.append(" ".join(current))

        return [c for c in chunks if len(c.strip()) > 50]
